"""Generate Haikal-Fitri-Resume.docx using python-docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

OUTPUT = Path(__file__).parent / "assets" / "Haikal-Fitri-Resume.docx"

PRIMARY = RGBColor(0x0F, 0x4C, 0x75)
ACCENT = RGBColor(0x32, 0x82, 0xB8)
MUTED = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0, 0, 0)


def set_run(run, *, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", *, size=10.5, bold=False, italic=False,
                  color=BLACK, align=None, space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bottom_border(paragraph, color_hex="3282B8", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run(run, size=11.5, bold=True, color=PRIMARY)
    add_bottom_border(p)
    return p


def add_bullet(doc, runs):
    """runs is a list of (text, {bold, italic, color, size}) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.6)
    for text, opts in runs:
        run = p.add_run(text)
        set_run(
            run,
            size=opts.get("size", 10),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
            color=opts.get("color", BLACK),
        )
    return p


def add_bullet_text(doc, text):
    return add_bullet(doc, [(text, {"size": 10})])


def add_role(doc, title, meta):
    p_title = add_paragraph(doc, title, size=10.5, bold=True, color=BLACK, space_after=0)
    p_meta = add_paragraph(doc, meta, size=10, italic=True, color=MUTED, space_after=3)
    return p_title, p_meta


def add_skill_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    set_run(run_label, size=10, bold=True, color=PRIMARY)
    run_value = p.add_run(value)
    set_run(run_value, size=10, color=BLACK)


def build(doc):
    # Margins
    for section in doc.sections:
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)

    # Header
    add_paragraph(
        doc,
        "Muhammad Haikal Fitri Bin Hassim @ Husain",
        size=20,
        bold=True,
        color=PRIMARY,
        space_after=2,
    )
    add_paragraph(
        doc,
        "Engineer, Digital Infrastructure Support",
        size=11.5,
        color=ACCENT,
        space_after=4,
    )
    add_paragraph(
        doc,
        "+60 111 932 9198  |  haikalfitri.hassim@gmail.com  |  linkedin.com/in/muhdhaikalfitri",
        size=9.5,
        color=MUTED,
        space_after=2,
    )
    header_loc = add_paragraph(
        doc,
        "Subang Jaya / Kajang, Selangor  |  haikalfitri.github.io",
        size=9.5,
        color=MUTED,
        space_after=4,
    )
    add_bottom_border(header_loc)

    # Experience
    add_section_heading(doc, "Experience")

    add_role(doc, "Engineer, Digital Infrastructure Support", "Cypark Sdn Bhd  ·  2025 – Present")
    for item in [
        "Manages Proxmox virtualization environment and backup systems to ensure platform reliability.",
        "Conducts SCADA vulnerability assessments to safeguard operational technology networks.",
        "Manages firewalls, switches, and access points across the corporate network.",
        "Deployed Splashtop remote access to support distributed engineering and operations teams.",
        "Handles IT onboarding and end-user device provisioning across the organisation.",
        "Administers Microsoft 365 and SharePoint, including identity, licensing, and collaboration workspaces.",
        "Drives AI-driven workflow automation to streamline internal IT and business processes.",
    ]:
        add_bullet_text(doc, item)

    add_role(doc, "Business / System Analyst (GEES)", "PETRONAS Carigali Sdn Bhd  ·  Jul 2023 – Jul 2024")
    for item in [
        "Led data integration for the Marine Intelligence System (MARSIS) to unify logistics data sources.",
        "Played a key role in User Acceptance Testing (UAT) of digital logistics tools prior to production rollout.",
        "Streamlined logistic digital tools, materially reducing manhours required for routine reporting.",
        "Implemented the HSE Assurance Program tracking workflow used across operational teams.",
        "Coordinated the Mentorship Program and a fuel benchmarking initiative applying Lean Six Sigma methods.",
        "Facilitated engagement and onboarding with over 40 contractors and external partners.",
    ]:
        add_bullet_text(doc, item)

    add_role(doc, "Web Developer Intern", "UiTM Technoventure  ·  Sep 2022 – Jan 2023")
    for item in [
        "Conducted stakeholder interviews to gather and document system requirements.",
        "Developed flowcharts and process diagrams describing core system functionalities.",
        "Managed MySQL database operations including schema design, queries, and maintenance.",
        "Enhanced UI/UX design across internal web applications.",
        "Engineered web pages and forms using HTML, CSS, and PHP.",
    ]:
        add_bullet_text(doc, item)

    # Education
    add_section_heading(doc, "Education")
    add_role(
        doc,
        "Bachelor of Information Systems (Hons.) — Intelligent Systems Engineering",
        "Universiti Teknologi MARA (UiTM)  ·  2019 – 2023  ·  CGPA 3.59 (First Class)",
    )
    add_role(
        doc,
        "Malaysian Matriculation Programme",
        "Kolej Matrikulasi Kelantan  ·  2018 – 2019  ·  CGPA 3.13",
    )

    # Technical Skills
    add_section_heading(doc, "Technical Skills")
    add_skill_line(
        doc,
        "Systems & IT",
        "Proxmox, Network Configuration, Network Security, Remote Access (Splashtop), "
        "Microsoft 365 Admin, SharePoint, SCADA Networks, AI Workflow Automation",
    )
    add_skill_line(doc, "Tools", "Wireshark, pfSense / Firewall, Power BI, Tableau, Excel, SQL Server, phpMyAdmin")
    add_skill_line(doc, "Programming", "Python, SQL, R, Java, HTML, CSS, JavaScript")
    add_skill_line(doc, "Data Science", "Data Analysis, Machine Learning, NLP, Data Visualization")
    add_skill_line(
        doc,
        "Software Development",
        "Version Control, UAT, Data Integration, Database Management, Agile",
    )
    add_skill_line(
        doc,
        "Soft Skills",
        "Problem Solving, Adaptability, Collaboration, Integrity, Resilience, "
        "Continuous Learning, Attention to Detail",
    )

    # Personal Projects
    add_section_heading(doc, "Personal Projects")
    add_bullet(
        doc,
        [
            ("Twitter Sentiment Analysis of Open Distance Learning (ODL)", {"size": 10, "bold": True}),
            (" — NLP, Pandas, RapidMiner, Power BI.", {"size": 10}),
        ],
    )
    add_bullet(
        doc,
        [
            ("Monster Movie Analysis App with Predictive Modeling", {"size": 10, "bold": True}),
            (" — Shiny web app with predictive modelling.", {"size": 10}),
        ],
    )
    add_bullet(
        doc,
        [
            ("FM24: Wonderkid Recruiting Analysis", {"size": 10, "bold": True}),
            (" — Statistical analysis and interactive dashboards.", {"size": 10}),
        ],
    )

    # Certifications
    add_section_heading(doc, "Certifications")
    for item in [
        "2024  ·  PETRONAS Lean Six Sigma Yellow Belt.",
        "2024  ·  PETRONAS Citizen Analytics Certificate.",
        "2023  ·  Coursera — Foundations: Data, Data, Everywhere.",
    ]:
        add_bullet_text(doc, item)

    # References
    add_section_heading(doc, "References")
    add_bullet(
        doc,
        [
            ("Loga Nathan Suppiah", {"size": 10, "bold": True}),
            (" — Manager, PETRONAS Carigali  ·  loganathan@petronas.com.my", {"size": 10}),
        ],
    )
    add_bullet(
        doc,
        [
            ("Arthur L Arulananda Savarydass", {"size": 10, "bold": True}),
            (" — Specialist, PETRONAS Carigali  ·  arthur.savarydass@petronas.com.my", {"size": 10}),
        ],
    )


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    build(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
